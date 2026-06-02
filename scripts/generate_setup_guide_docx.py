"""Generate docs/SETUP_GUIDE.docx from the SETUP_GUIDE.md content."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = Path(__file__).parent.parent / "docs" / "SETUP_GUIDE.docx"

# ── Colour palette ──────────────────────────────────────────────────────────
BLUE_DARK   = RGBColor(0x1E, 0x3A, 0x5F)   # headings
BLUE_MID    = RGBColor(0x2E, 0x6D, 0xA4)   # h2
BLUE_LIGHT  = RGBColor(0x35, 0x86, 0xC8)   # h3
ACCENT      = RGBColor(0xE8, 0x54, 0x7A)   # highlight / role tags
CODE_BG     = RGBColor(0xF3, 0xF4, 0xF6)   # code block background
CODE_FG     = RGBColor(0x1F, 0x2D, 0x3D)   # code text
BODY        = RGBColor(0x1A, 0x1A, 0x2E)   # normal text
TABLE_HEAD  = RGBColor(0x1E, 0x3A, 0x5F)
TABLE_ALT   = RGBColor(0xEB, 0xF2, 0xFA)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(22)
        run.font.color.rgb = BLUE_DARK
    elif level == 2:
        run.font.size  = Pt(16)
        run.font.color.rgb = BLUE_MID
    else:
        run.font.size  = Pt(13)
        run.font.color.rgb = BLUE_LIGHT
    run.font.name = "Calibri"
    return p


def add_body(doc: Document, text: str, bold_parts: list[str] | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_parts:
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx >= 0:
                before = remaining[:idx]
                if before:
                    r = p.add_run(before)
                    r.font.color.rgb = BODY
                    r.font.name = "Calibri"
                    r.font.size = Pt(11)
                r = p.add_run(bp)
                r.bold = True
                r.font.color.rgb = BODY
                r.font.name = "Calibri"
                r.font.size = Pt(11)
                remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.color.rgb = BODY
            r.font.name = "Calibri"
            r.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.color.rgb = BODY
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def add_note(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x66, 0x77)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    return p


def add_code(doc: Document, lines: str):
    for line in lines.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F3F4F6")
        pPr.append(shd)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)
        run.font.color.rgb = CODE_FG


def add_bullet(doc: Document, text: str, sub: bool = False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(1.0 if not sub else 1.8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.color.rgb = BODY
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "1E3A5F")
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "EBF2FA" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.clear()
            # Handle backtick-wrapped code inline
            if val.startswith("`") and val.endswith("`"):
                run = p.add_run(val[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(9.5)
                run.font.color.rgb = ACCENT
            else:
                run = p.add_run(val)
                run.font.name = "Calibri"
                run.font.size = Pt(10.5)
                run.font.color.rgb = BODY
    doc.add_paragraph()


def add_divider(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E6DA4")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Build the document ───────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Cover / Title ────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run("AI Interaction Tracker")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = BLUE_DARK
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Simple Setup & User Guide")
run.font.size = Pt(16)
run.font.color.rgb = BLUE_MID
run.font.name = "Calibri"
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Phase 1 — Python File Parser for AI IDE Tools")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x66, 0x77)
run.font.name = "Calibri"

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

# ── What This Tool Does ──────────────────────────────────────────────────────
add_heading(doc, "What This Tool Does", 2)
add_body(doc,
    "Every time you chat with an AI in your IDE, the conversation is saved to a log file on your disk. "
    "This tool reads those log files and exports them into a clean, structured CSV file you can open in Excel.",
    bold_parts=["Claude Code", "Antigravity", "Codex"])
add_note(doc, "No internet. No API keys. No installation of extra software.")
add_divider(doc)

# ── Requirements ─────────────────────────────────────────────────────────────
add_heading(doc, "Requirements", 2)
add_table(doc,
    ["Requirement", "Details"],
    [
        ["Python",           "3.10 or higher"],
        ["Operating System", "Windows 10 / 11"],
        ["AI IDE",           "Claude Code, Antigravity, or Codex (at least one)"],
    ]
)
add_divider(doc)

# ── Step 1 — Install ─────────────────────────────────────────────────────────
add_heading(doc, "Step 1 — Install", 2)
add_body(doc, "Open PowerShell and run these two commands once:")
add_code(doc, """cd "C:\\Users\\Robin\\AI TRACKING SYSTEM PYTHON SCRIPT"
.venv\\Scripts\\pip.exe install -e .""")

add_body(doc, "Verify the install worked:")
add_code(doc, "ai-tracker list-tools")

add_body(doc, "Expected output:")
add_code(doc, """Tool             Status       Path
----------------------------------------------------------------------
  antigravity    [found    ]  C:\\Users\\Robin\\.gemini\\antigravity-ide\\brain
  claudecode     [found    ]  C:\\Users\\Robin\\.claude\\projects
  codex          [found    ]  C:\\Users\\Robin\\.codex""")

add_note(doc, "[found] means the tool is installed and has conversation data ready to export.")
add_divider(doc)

# ── Step 2 — First Export ────────────────────────────────────────────────────
add_heading(doc, "Step 2 — Your First Export", 2)
add_body(doc, "Export everything from all tools into one CSV:")
add_code(doc, "ai-tracker parse --tool all -o my_interactions.csv")

add_body(doc, "What you will see:")
add_code(doc, """  [antigravity] 477 message(s) parsed.
  [claudecode]  821 message(s) parsed.
  [codex]         0 message(s) parsed.

Exported 1298 message(s) -> my_interactions.csv""")

add_note(doc, "Open my_interactions.csv in Excel — that is it.")
add_divider(doc)

# ── Step 3 — CSV Columns ─────────────────────────────────────────────────────
add_heading(doc, "Step 3 — Understand the CSV", 2)
add_body(doc, "Each row is one message — either something you typed or the AI's reply.")
add_table(doc,
    ["Column", "What it means", "Example"],
    [
        ["`project`",    "Which project the conversation belongs to", "Sparq"],
        ["`session_id`", "Unique ID for that conversation",           "b65bda2b-bc85-..."],
        ["`timestamp`",  "Exact date and time",                       "2026-05-18T09:22:28"],
        ["`role`",       "Who spoke",                                 "human / assistant / tool"],
        ["`message`",    "The full text — never cut short",           "How do I sort a dict?"],
        ["`tool`",       "Which AI IDE was used",                     "claudecode / antigravity"],
        ["`file_path`",  "Where the log file lives on disk",          "C:\\Users\\Robin\\.claude\\..."],
    ]
)

add_heading(doc, "What each role means", 3)
add_bullet(doc, "human     — A prompt you typed")
add_bullet(doc, "assistant — The AI's response")
add_bullet(doc, "tool      — A background action the AI took (reading a file, listing a folder, running a command)")
add_divider(doc)

# ── All Commands ─────────────────────────────────────────────────────────────
add_heading(doc, "All Available Commands", 2)

add_heading(doc, "Check which tools have data", 3)
add_code(doc, "ai-tracker list-tools")

add_heading(doc, "Export everything", 3)
add_code(doc, "ai-tracker parse --tool all -o output.csv")

add_heading(doc, "Export one tool only", 3)
add_code(doc, """ai-tracker parse --tool claudecode  -o claude.csv
ai-tracker parse --tool antigravity -o antigravity.csv
ai-tracker parse --tool codex       -o codex.csv""")

add_heading(doc, "Filter by date", 3)
add_code(doc, """# Today only
ai-tracker parse --tool all --start-date 2026-06-01 --end-date 2026-06-01 -o today.csv

# A full month
ai-tracker parse --tool all --start-date 2026-05-01 --end-date 2026-05-31 -o may.csv

# From a specific date onwards
ai-tracker parse --tool all --start-date 2026-05-25 -o recent.csv""")

add_heading(doc, "Filter by project name", 3)
add_code(doc, """# Exact project
ai-tracker parse --tool all --project "Sparq" -o sparq.csv

# Partial match (case-insensitive)
ai-tracker parse --tool all --project "tracking" -o tracking.csv""")

add_heading(doc, "Split into one file per project", 3)
add_body(doc, "This creates a folder with one CSV per project — ideal for auditing individual projects.")
add_code(doc, "ai-tracker parse --tool all --split-by-project -o projects/")
add_body(doc, "Output folder:")
add_code(doc, """projects/
  sparq.csv
  ai_tracking_system_python_script.csv
  ai_interaction_tracking_system.csv
  ai_prompt_tracker.csv
  general.csv""")

add_heading(doc, "Combine multiple filters", 3)
add_code(doc, """ai-tracker parse --tool all --project "Sparq" --start-date 2026-05-01 --end-date 2026-05-31 --split-by-project -o sparq_may/""")

add_heading(doc, "Exclude Claude Code subagent threads", 3)
add_code(doc, "ai-tracker parse --tool claudecode --no-sidechains -o main_only.csv")
add_divider(doc)

# ── Session IDs ───────────────────────────────────────────────────────────────
add_heading(doc, "How Session IDs Work", 2)
add_body(doc,
    "Each conversation gets a unique ID that comes directly from the IDE. "
    "The tracker never makes one up.")
add_table(doc,
    ["Tool", "Where the session ID comes from"],
    [
        ["Claude Code", "The .jsonl filename under ~/.claude/projects/"],
        ["Antigravity", "The UUID folder name under ~/.gemini/antigravity-ide/brain/"],
        ["Codex",       "A field inside the session file"],
    ]
)
add_note(doc, "Re-running the export tomorrow gives the same session IDs. The CSV is stable and consistent over time.")
add_divider(doc)

# ── Project Detection ─────────────────────────────────────────────────────────
add_heading(doc, "How Projects Are Detected", 2)
add_body(doc,
    "The tool automatically figures out which project each conversation belongs to.")
add_table(doc,
    ["Tool", "How it finds the project name"],
    [
        ["Claude Code", "Reads the project folder name from the file path"],
        ["Antigravity", "Scans the transcript for Active Document: or workspace paths in the metadata"],
        ["Codex",       "Reads the project folder name from the file path"],
    ]
)
add_note(doc, 'Raw folder names like "c--Users-Robin-AI-TRACKING-SYSTEM-PYTHON-SCRIPT" are automatically cleaned to "Ai Tracking System Python Script".')
add_divider(doc)

# ── Latency ───────────────────────────────────────────────────────────────────
add_heading(doc, "Latency — How Fresh is the Data?", 2)
add_table(doc,
    ["Tool", "Delay after you send a message"],
    [
        ["Claude Code", "~0.2 seconds"],
        ["Antigravity", "~2–3 seconds"],
    ]
)
add_note(doc,
    "The parser reads what is on disk at the moment you run the command. "
    "The currently active message (being typed right now) will not appear until it is complete and flushed to disk.")
add_divider(doc)

# ── Troubleshooting ───────────────────────────────────────────────────────────
add_heading(doc, "Troubleshooting", 2)

add_heading(doc, "ai-tracker not recognized as a command", 3)
add_body(doc, "Run with the full venv path instead:")
add_code(doc, """cd "C:\\Users\\Robin\\AI TRACKING SYSTEM PYTHON SCRIPT"
.venv\\Scripts\\python.exe -m ai_tracker.cli parse --tool all -o output.csv""")

add_heading(doc, "Tool shows [not found]", 3)
add_body(doc,
    "The tool is either not installed or has never been used. "
    "Start a conversation in that IDE first, then re-run ai-tracker list-tools.")

add_heading(doc, "No messages found", 3)
add_body(doc,
    "Check the date filter. The --end-date covers the full day automatically. "
    "Also verify the tool path with list-tools.")

add_heading(doc, "CSV opens with garbled characters in Excel", 3)
add_body(doc, "Open Excel > Data > From Text/CSV > select the file > choose UTF-8 encoding.")

add_heading(doc, "Run tests to check everything is working", 3)
add_code(doc, ".venv\\Scripts\\pytest.exe tests/ -v")
add_note(doc, "Expected result: 128 passed")
add_divider(doc)

# ── Quick Reference Card ──────────────────────────────────────────────────────
add_heading(doc, "Quick Reference Card", 2)
add_code(doc, """SETUP
  cd "C:\\Users\\Robin\\AI TRACKING SYSTEM PYTHON SCRIPT"
  .venv\\Scripts\\pip.exe install -e .

CHECK
  ai-tracker list-tools

EXPORT
  ai-tracker parse --tool all -o output.csv
  ai-tracker parse --tool claudecode  -o claude.csv
  ai-tracker parse --tool antigravity -o antigravity.csv

FILTER BY DATE
  ai-tracker parse --tool all --start-date 2026-06-01 -o today.csv
  ai-tracker parse --tool all --start-date 2026-05-01 --end-date 2026-05-31 -o may.csv

FILTER BY PROJECT
  ai-tracker parse --tool all --project "Sparq" -o sparq.csv

SPLIT BY PROJECT
  ai-tracker parse --tool all --split-by-project -o projects/""")

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
